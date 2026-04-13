# -*- coding: utf-8 -*-
"""
Generador de Informe Ejecutivo - Ecosistema Digital Frutos Tropicales Peru S.A.C.
Autor: Script generado para Rodrigo Garcia
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy

# ── Colors ──
GREEN_DARK = RGBColor(0x0d, 0x3d, 0x15)
GREEN_PRIMARY = RGBColor(0x16, 0xa3, 0x4a)
GREEN_LIGHT = RGBColor(0xdc, 0xfc, 0xe7)
GREEN_ACCENT = RGBColor(0x22, 0xc5, 0x5e)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x1a, 0x1a, 0x1a)
GRAY = RGBColor(0x4a, 0x4a, 0x4a)
GRAY_LIGHT = RGBColor(0x6b, 0x6b, 0x6b)
GRAY_BG = RGBColor(0xf8, 0xf9, 0xfa)

OUTPUT_PATH = r"C:\Users\ARES\Desktop\FRUTOS TROPICALES\PAGINA WEB- FRUTOS TROPICALES\Informe_Ejecutivo_Ecosistema_Digital_FTP.docx"

doc = Document()

# ── Page Setup ──
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = BLACK

# ── Helper Functions ──

# OOXML requires pPr children in a specific order. pBdr must come after
# pStyle, keepNext, keepLines, pageBreakBefore, framePr, widowControl, numPr,
# suppressLineNumbers, but BEFORE tabs, suppressAutoHyphens, kinsoku, etc.
# The safest approach: insert pBdr right after any existing early elements.
PPR_ORDER = [
    'w:pStyle', 'w:keepNext', 'w:keepLines', 'w:pageBreakBefore',
    'w:framePr', 'w:widowControl', 'w:numPr', 'w:suppressLineNumbers',
    'w:pBdr', 'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku',
    'w:wordWrap', 'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE',
    'w:autoSpaceDN', 'w:bidi', 'w:adjustRightInd', 'w:snapToGrid',
    'w:spacing', 'w:ind', 'w:contextualSpacing', 'w:mirrorIndents',
    'w:suppressOverlap', 'w:jc', 'w:textDirection', 'w:textAlignment',
    'w:textboxTightWrap', 'w:outlineLvl', 'w:divId', 'w:cnfStyle',
    'w:rPr', 'w:sectPr', 'w:pPrChange'
]

def _get_ppr_order_index(tag_name):
    """Get the order index for a pPr child element."""
    # Extract local name from full qualified name
    if '}' in tag_name:
        local = 'w:' + tag_name.split('}')[-1]
    else:
        local = tag_name
    try:
        return PPR_ORDER.index(local)
    except ValueError:
        return 999

def _insert_pBdr(pPr, pBdr_elem):
    """Insert pBdr element at the correct position in pPr."""
    target_idx = _get_ppr_order_index(pBdr_elem.tag)
    insert_before = None
    for child in pPr:
        child_idx = _get_ppr_order_index(child.tag)
        if child_idx > target_idx:
            insert_before = child
            break
    if insert_before is not None:
        pPr.insert(list(pPr).index(insert_before), pBdr_elem)
    else:
        pPr.append(pBdr_elem)

def _insert_pPr_child(pPr, elem):
    """Insert any pPr child element at the correct position."""
    _insert_pBdr(pPr, elem)  # Same logic works for any pPr child

TCPR_ORDER = [
    'w:cnfStyle', 'w:tcW', 'w:gridSpan', 'w:hMerge', 'w:vMerge',
    'w:tcBorders', 'w:shd', 'w:noWrap', 'w:tcMar', 'w:textDirection',
    'w:tcFitText', 'w:vAlign', 'w:hideMark', 'w:headers',
    'w:cellIns', 'w:cellDel', 'w:cellMerge', 'w:tcPrChange'
]

def _get_tcpr_order_index(tag_name):
    if '}' in tag_name:
        local = 'w:' + tag_name.split('}')[-1]
    else:
        local = tag_name
    try:
        return TCPR_ORDER.index(local)
    except ValueError:
        return 999

def _insert_tcPr_child(tcPr, elem):
    """Insert element at the correct position in tcPr."""
    target_idx = _get_tcpr_order_index(elem.tag)
    insert_before = None
    for child in tcPr:
        child_idx = _get_tcpr_order_index(child.tag)
        if child_idx > target_idx:
            insert_before = child
            break
    if insert_before is not None:
        tcPr.insert(list(tcPr).index(insert_before), elem)
    else:
        tcPr.append(elem)

def set_cell_shading(cell, color_hex):
    """Set background color for a table cell."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{color_hex}"/>')
    tcPr = cell._tc.get_or_add_tcPr()
    # Remove existing shading
    existing = tcPr.findall(qn('w:shd'))
    for e in existing:
        tcPr.remove(e)
    _insert_tcPr_child(tcPr, shading_elm)

TC_BORDER_ORDER = ['top', 'start', 'left', 'bottom', 'end', 'right', 'insideH', 'insideV', 'tl2br', 'tr2bl']

def set_cell_border(cell, **kwargs):
    """Set borders on a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Remove existing borders
    existing = tcPr.findall(qn('w:tcBorders'))
    for e in existing:
        tcPr.remove(e)
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    # Insert border elements in schema-required order
    for edge in TC_BORDER_ORDER:
        if edge in kwargs:
            val = kwargs[edge]
            element = parse_xml(
                f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
                f'w:sz="{val.get("sz", "4")}" w:space="0" '
                f'w:color="{val.get("color", "000000")}"/>'
            )
            tcBorders.append(element)
    _insert_tcPr_child(tcPr, tcBorders)

def add_formatted_paragraph(text, bold=False, size=11, color=BLACK, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6, font_name='Calibri', italic=False):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = font_name
    return p

def add_section_header(number, title):
    """Add a numbered section header with green styling."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    # Add bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="8" w:space="4" w:color="16a34a"/>'
        f'</w:pBdr>'
    )
    # Insert pBdr at the correct position in pPr
    _insert_pBdr(pPr, pBdr)

    run = p.add_run(f"{number}. {title}")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = GREEN_DARK
    run.font.name = 'Calibri'
    return p

def add_subsection_header(title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = GREEN_PRIMARY
    run.font.name = 'Calibri'
    return p

def add_bullet(text, bold_prefix="", level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(1.5 + level * 0.8)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = BLACK
        run.font.name = 'Calibri'
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = GRAY
    run.font.name = 'Calibri'
    return p

def create_styled_table(headers, rows, col_widths=None):
    """Create a professional table with green header."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Header row
    hdr_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = WHITE
        run.font.name = 'Calibri'
        set_cell_shading(cell, "16a34a")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9)
            run.font.color.rgb = BLACK
            run.font.name = 'Calibri'
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # Alternate row shading
            if r_idx % 2 == 1:
                set_cell_shading(cell, "f0fdf4")

    # Set borders for all cells
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell,
                top={"sz": "4", "color": "d1d5db", "val": "single"},
                bottom={"sz": "4", "color": "d1d5db", "val": "single"},
                left={"sz": "4", "color": "d1d5db", "val": "single"},
                right={"sz": "4", "color": "d1d5db", "val": "single"},
            )

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    return table

def add_kpi_box(label, value, description=""):
    """Add a KPI-style paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("  \u25CF  ")
    run.font.color.rgb = GREEN_PRIMARY
    run.font.size = Pt(11)
    run = p.add_run(f"{label}: ")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = GREEN_DARK
    run = p.add_run(value)
    run.font.size = Pt(11)
    run.font.color.rgb = BLACK
    if description:
        run = p.add_run(f" \u2014 {description}")
        run.font.size = Pt(9.5)
        run.font.color.rgb = GRAY_LIGHT

def add_page_break():
    doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════

# Spacer
for _ in range(4):
    add_formatted_paragraph("", size=12)

# Top green line
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
pPr = p._p.get_or_add_pPr()
pBdr = parse_xml(
    f'<w:pBdr {nsdecls("w")}>'
    f'<w:top w:val="single" w:sz="36" w:space="1" w:color="16a34a"/>'
    f'</w:pBdr>'
)
_insert_pPr_child(pPr, pBdr)

# Company name
add_formatted_paragraph("FRUTOS TROPICALES PERU S.A.C.", bold=True, size=28,
                        color=GREEN_DARK, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                        space_before=12, space_after=6)

# Thin separator
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
pPr = p._p.get_or_add_pPr()
pBdr = parse_xml(
    f'<w:pBdr {nsdecls("w")}>'
    f'<w:bottom w:val="single" w:sz="12" w:space="1" w:color="22c55e"/>'
    f'</w:pBdr>'
)
_insert_pPr_child(pPr, pBdr)

add_formatted_paragraph("", size=6)

# Title
add_formatted_paragraph("INFORME EJECUTIVO", bold=True, size=22,
                        color=GREEN_PRIMARY, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                        space_before=10, space_after=4)

add_formatted_paragraph("Ecosistema Digital de Gesti\u00f3n Integral", bold=True, size=18,
                        color=GREEN_DARK, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                        space_before=0, space_after=4)

add_formatted_paragraph("", size=6)

# Subtitle
add_formatted_paragraph("Presentaci\u00f3n a Gerencia General", bold=False, size=14,
                        color=GRAY, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                        space_before=0, space_after=24, italic=True)

# Bottom green line
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
pPr = p._p.get_or_add_pPr()
pBdr = parse_xml(
    f'<w:pBdr {nsdecls("w")}>'
    f'<w:bottom w:val="single" w:sz="24" w:space="1" w:color="16a34a"/>'
    f'</w:pBdr>'
)
_insert_pPr_child(pPr, pBdr)

# Spacer
for _ in range(3):
    add_formatted_paragraph("", size=10)

# Author info
add_formatted_paragraph("Elaborado por:", bold=False, size=11,
                        color=GRAY_LIGHT, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                        space_before=0, space_after=2)

add_formatted_paragraph("Rodrigo Garcia", bold=True, size=16,
                        color=GREEN_DARK, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                        space_before=0, space_after=2)

add_formatted_paragraph("Systems Analyst & Developer-QA", bold=False, size=12,
                        color=GREEN_PRIMARY, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                        space_before=0, space_after=2, italic=True)

add_formatted_paragraph("Frutos Tropicales Peru S.A.C.", bold=False, size=11,
                        color=GRAY, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                        space_before=0, space_after=24)

# Date
add_formatted_paragraph("20 de marzo de 2026", bold=True, size=13,
                        color=GREEN_DARK, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                        space_before=0, space_after=6)

# Confidential note
add_formatted_paragraph("DOCUMENTO CONFIDENCIAL", bold=True, size=9,
                        color=RGBColor(0xdc, 0x26, 0x26), alignment=WD_ALIGN_PARAGRAPH.CENTER,
                        space_before=12, space_after=0)

add_page_break()

# ══════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (styled manually)
# ══════════════════════════════════════════════════════════════════

add_formatted_paragraph("\u00cdNDICE", bold=True, size=18, color=GREEN_DARK,
                        alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=12, space_after=16)

toc_items = [
    ("1.", "Resumen Ejecutivo"),
    ("2.", "\u00bfPor qu\u00e9 es importante contar con estas herramientas?"),
    ("3.", "Ecosistema Digital \u2014 Componentes Desarrollados"),
    ("4.", "Herramientas y Flujo de Trabajo por \u00c1rea"),
    ("5.", "Beneficios Clave para la Empresa"),
    ("6.", "Mejoras Planificadas a Corto Plazo"),
    ("7.", "Visi\u00f3n a Futuro \u2014 Planta de Frescos"),
    ("8.", "Visi\u00f3n a Futuro \u2014 Innovaciones Tecnol\u00f3gicas"),
    ("9.", "Justificaci\u00f3n de la Inversi\u00f3n \u2014 Plan MAX"),
    ("10.", "Conclusi\u00f3n"),
]

for num, title in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(1.2))
    run = p.add_run(f"{num}")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = GREEN_PRIMARY
    run = p.add_run(f"\t{title}")
    run.font.size = Pt(11)
    run.font.color.rgb = BLACK

add_page_break()

# ══════════════════════════════════════════════════════════════════
# SECTION 1: RESUMEN EJECUTIVO
# ══════════════════════════════════════════════════════════════════

add_section_header("1", "Resumen Ejecutivo")

add_formatted_paragraph(
    "Frutos Tropicales Peru S.A.C. ha dado un paso decisivo hacia la transformaci\u00f3n digital "
    "mediante el desarrollo de un ecosistema integral de herramientas tecnol\u00f3gicas dise\u00f1adas "
    "espec\u00edficamente para las necesidades operativas, de calidad y gerenciales de nuestra empresa. "
    "Este ecosistema abarca desde la presencia corporativa en l\u00ednea hasta dashboards interactivos "
    "de gesti\u00f3n en tiempo real.",
    size=11, color=BLACK, space_before=6, space_after=8
)

add_formatted_paragraph(
    "Todo este ecosistema ha sido concebido, dise\u00f1ado, desarrollado e implementado por una sola persona: "
    "Rodrigo Garcia, Systems Analyst & Developer-QA, utilizando tecnolog\u00edas web modernas, bases de datos "
    "en la nube y frameworks de visualizaci\u00f3n de datos de \u00faltima generaci\u00f3n.",
    bold=True, size=11, color=GREEN_DARK, space_before=4, space_after=8
)

add_formatted_paragraph(
    "El ecosistema comprende m\u00e1s de 15 componentes funcionales que cubren las \u00e1reas de producci\u00f3n, "
    "calidad, laboratorio, almac\u00e9n, log\u00edstica y gerencia. Cada componente proporciona informaci\u00f3n "
    "cr\u00edtica para la toma de decisiones, eliminando la dependencia de procesos manuales, hojas de c\u00e1lculo "
    "dispersas y comunicaciones informales.",
    size=11, color=BLACK, space_before=4, space_after=8
)

add_subsection_header("Cifras clave del ecosistema:")

kpis = [
    ("Componentes desarrollados", "15+ herramientas y dashboards"),
    ("P\u00e1gina web corporativa", "1 sitio responsive con cat\u00e1logo, certificaciones y formulario de contacto"),
    ("Portal de gerencia", "1 portal con login seguro, roles diferenciados y m\u00faltiples paneles"),
    ("Dashboards interactivos", "8+ dashboards con filtros, gr\u00e1ficos y datos en tiempo real"),
    ("Certificaciones cubiertas", "BRC, HACCP, Org\u00e1nico UE, FDA, GlobalG.A.P."),
    ("Almacenes monitoreados", "3 almacenes (FTP, Agroempaques, Ransa)"),
    ("Stock controlado", "272,020 kg / 24,210 cajas / 10.10 FCL"),
    ("Alertas activas", "7 alertas de temperatura en tiempo real"),
    ("Productos de laboratorio", "30+ \u00edtems con sem\u00e1foro de abastecimiento"),
]

for label, value in kpis:
    add_kpi_box(label, value)

add_page_break()

# ══════════════════════════════════════════════════════════════════
# SECTION 2: POR QUE ES IMPORTANTE
# ══════════════════════════════════════════════════════════════════

add_section_header("2", "\u00bfPor qu\u00e9 es importante contar con estas herramientas?")

add_formatted_paragraph(
    "En el competitivo mercado de exportaci\u00f3n de frutas tropicales congeladas IQF, la diferencia entre "
    "una empresa l\u00edder y una rezagada radica en su capacidad para gestionar informaci\u00f3n de manera "
    "eficiente y tomar decisiones basadas en datos. Las herramientas digitales desarrolladas abordan "
    "directamente los desaf\u00edos cr\u00edticos de nuestra operaci\u00f3n:",
    size=11, space_before=6, space_after=10
)

reasons = [
    ("Toma de decisiones en tiempo real: ",
     "La gerencia tiene acceso inmediato a indicadores de producci\u00f3n, stock, calidad y "
     "temperaturas sin depender de reportes manuales que llegan con horas o d\u00edas de retraso. "
     "Un gerente puede verificar el avance de producci\u00f3n del turno actual desde su celular."),
    ("Visibilidad total para la gerencia: ",
     "El portal de gerencia centraliza toda la informaci\u00f3n operativa en un solo punto de acceso "
     "con roles diferenciados. Cada nivel jer\u00e1rquico ve exactamente lo que necesita, desde el "
     "resumen ejecutivo hasta el detalle operativo por hora."),
    ("Trazabilidad y cumplimiento normativo: ",
     "Las certificaciones BRC, HACCP, FDA y Org\u00e1nico UE exigen trazabilidad rigurosa. "
     "Nuestros dashboards registran autom\u00e1ticamente temperaturas, consumos de laboratorio, "
     "resultados de an\u00e1lisis y par\u00e1metros de calidad, generando evidencia digital para auditor\u00edas."),
    ("Eficiencia operativa y reducci\u00f3n de mermas: ",
     "El seguimiento hora a hora de la producci\u00f3n (ej: Mango Kent ~2 tn/hora) permite identificar "
     "ca\u00eddas de rendimiento inmediatamente. El monitoreo de temperaturas con alertas previene "
     "p\u00e9rdidas por quiebres de cadena de fr\u00edo."),
    ("Ventaja competitiva en el mercado exportador: ",
     "Clientes internacionales exigen cada vez m\u00e1s transparencia y datos. Contar con dashboards "
     "profesionales y datos estructurados posiciona a Frutos Tropicales como una empresa moderna "
     "y confiable ante compradores de EE.UU., Europa y Asia."),
    ("Ahorro de costos vs. procesos manuales: ",
     "Cada hora que un supervisor dedica a llenar planillas en papel, consolidar datos en Excel "
     "y enviar correos con reportes, es una hora que podr\u00eda invertir en supervisar la operaci\u00f3n. "
     "La digitalizaci\u00f3n reduce errores humanos y elimina duplicidad de esfuerzos."),
]

for bold_part, text in reasons:
    add_bullet(text, bold_prefix=bold_part)

add_page_break()

# ══════════════════════════════════════════════════════════════════
# SECTION 3: ECOSISTEMA DIGITAL - COMPONENTES
# ══════════════════════════════════════════════════════════════════

add_section_header("3", "Ecosistema Digital \u2014 Componentes Desarrollados")

add_formatted_paragraph(
    "A continuaci\u00f3n se presenta el detalle de cada componente del ecosistema digital, organizado "
    "por \u00e1rea funcional. Cada herramienta ha sido desarrollada con tecnolog\u00edas web modernas "
    "(HTML5, CSS3, JavaScript, Chart.js, Supabase) y dise\u00f1o responsive para acceso desde "
    "cualquier dispositivo.",
    size=11, space_before=6, space_after=12
)

add_subsection_header("3.1 P\u00e1gina Web Corporativa")

add_formatted_paragraph(
    "Sitio web profesional que representa la imagen de Frutos Tropicales Peru ante clientes "
    "internacionales y socios comerciales. Incluye video hero de impacto, cat\u00e1logo completo de "
    "productos IQF (Fresas, Ar\u00e1ndanos, Mango Congelado, Pi\u00f1a Golden, Palta Congelada, Granada), "
    "secci\u00f3n de procesos, certificaciones visibles (BRC, HACCP, Org\u00e1nico UE, FDA, GlobalG.A.P.) "
    "y formulario de contacto. Dise\u00f1o premium, totalmente responsive.",
    size=10.5, space_before=4, space_after=10
)

add_subsection_header("3.2 Portal de Gerencia")

add_formatted_paragraph(
    "Sistema central con autenticaci\u00f3n segura y control de acceso por roles (Gerencia General, "
    "Jefe de Producci\u00f3n, Jefe de Calidad). Incluye persistencia de sesi\u00f3n, cierre autom\u00e1tico "
    "por inactividad, modo oscuro/claro y paneles de navegaci\u00f3n a todos los m\u00f3dulos del sistema. "
    "Punto de entrada unificado para toda la informaci\u00f3n gerencial.",
    size=10.5, space_before=4, space_after=10
)

add_subsection_header("3.3 Tabla de Componentes del Ecosistema")

components_headers = ["N\u00b0", "Componente", "Funci\u00f3n Principal", "Caracter\u00edsticas Clave", "Usuarios"]

components_rows = [
    ["1", "P\u00e1gina Web Corporativa",
     "Presencia digital, imagen corporativa, captaci\u00f3n de clientes",
     "Video hero, cat\u00e1logo 6 productos IQF, certificaciones, formulario contacto, responsive",
     "Clientes, Socios, P\u00fablico"],
    ["2", "Portal de Gerencia",
     "Acceso centralizado seguro a todos los m\u00f3dulos",
     "Login por roles, persistencia sesi\u00f3n, auto-logout, modo oscuro/claro",
     "Gerencia, Jefes de \u00c1rea"],
    ["3", "Dashboard Resumen General",
     "Vista ejecutiva de KPIs principales",
     "Indicadores consolidados de toda la operaci\u00f3n",
     "Gerencia General"],
    ["4", "Indicadores de Producci\u00f3n",
     "M\u00e9tricas y seguimiento de eficiencia productiva",
     "Gr\u00e1ficos de tendencia, comparativos por turno y l\u00ednea",
     "Gerencia, Jefe Producci\u00f3n"],
    ["5", "Producci\u00f3n del D\u00eda",
     "Tracking diario de procesamiento hora a hora",
     "Mango Kent ~2tn/h, gr\u00e1ficos acumulados, KPI: 14,090 kg, 45.7% yield",
     "Supervisor, Jefe Producci\u00f3n"],
    ["6", "Avance de Contenedores",
     "Seguimiento de carga para exportaci\u00f3n",
     "Progreso de llenado, estado por contenedor",
     "Log\u00edstica, Gerencia"],
    ["7", "Stock General (Portal)",
     "Panel resumen de inventario de PT",
     "272,020 kg, 24,210 cajas, 10.10 FCL, 3 almacenes",
     "Gerencia, Almac\u00e9n"],
    ["8", "Dashboard Stock General",
     "Dashboard interactivo premium de inventarios",
     "4 gr\u00e1ficos interactivos, filtros por producto/planta, tabla ordenable, modo oscuro",
     "Gerencia, Almac\u00e9n, Comercial"],
    ["9", "Control de Temperaturas",
     "Monitoreo en tiempo real v\u00eda Supabase",
     "7 alertas activas, sem\u00e1foro, conexi\u00f3n cloud en tiempo real",
     "Calidad, Almac\u00e9n, Gerencia"],
    ["10", "Dashboard Laboratorio",
     "An\u00e1lisis microbiol\u00f3gicos y sem\u00e1foro stock insumos",
     "30+ productos, estados CR\u00cdTICO/BAJO/MODERADO/OK, filtro por unidad",
     "Jefe Calidad, Laboratorio"],
    ["11", "Consumo Laboratorio",
     "Tracking de consumos de insumos de laboratorio",
     "Gr\u00e1fico sem\u00e1foro, seguimiento por per\u00edodo",
     "Laboratorio, Calidad"],
    ["12", "Consumos de Calidad",
     "Seguimiento de suministros del \u00e1rea de calidad",
     "Dashboard de consumos, tendencias, alertas de stock",
     "Jefe Calidad"],
    ["13", "Costos de An\u00e1lisis Lab.",
     "Control de costos de an\u00e1lisis de laboratorio",
     "Seguimiento de gastos, comparativos por per\u00edodo",
     "Calidad, Finanzas"],
    ["14", "Certificaciones",
     "Gesti\u00f3n de certificaciones vigentes",
     "Estado de vigencia, renovaciones, documentaci\u00f3n",
     "Calidad, Gerencia"],
    ["15", "Reportes de Inspecci\u00f3n",
     "Gesti\u00f3n de reportes de inspecci\u00f3n de calidad",
     "Registro, seguimiento y archivo de inspecciones",
     "Calidad, Producci\u00f3n"],
]

create_styled_table(components_headers, components_rows)

add_page_break()

# ══════════════════════════════════════════════════════════════════
# SECTION 4: HERRAMIENTAS Y FLUJO DE TRABAJO POR AREA
# ══════════════════════════════════════════════════════════════════

add_section_header("4", "Herramientas y Flujo de Trabajo por \u00c1rea")

add_formatted_paragraph(
    "El siguiente cuadro muestra c\u00f3mo cada \u00e1rea de la empresa interact\u00faa con el ecosistema digital, "
    "detallando los roles, herramientas utilizadas, datos que ingresan, frecuencia de uso y "
    "dispositivos desde los cuales acceden.",
    size=11, space_before=6, space_after=12
)

workflow_headers = ["\u00c1rea", "Rol / Cargo", "Herramienta", "Datos que Ingresa", "Frecuencia", "Dispositivo"]

workflow_rows = [
    ["Producci\u00f3n", "Supervisor de L\u00ednea",
     "Producci\u00f3n del D\u00eda",
     "Kg procesados por hora, cajas producidas, rendimiento, incidencias",
     "Cada hora / Por turno",
     "Tablet / PC"],
    ["Producci\u00f3n", "Jefe de Producci\u00f3n",
     "Indicadores de Producci\u00f3n, Resumen General",
     "Metas diarias, validaci\u00f3n de datos, observaciones",
     "Diaria / Por turno",
     "PC / Tablet"],
    ["Producci\u00f3n", "Operador de L\u00ednea",
     "Producci\u00f3n del D\u00eda (visualizaci\u00f3n)",
     "Registro de paradas, peso de bandejas",
     "Continua",
     "Tablet"],
    ["Calidad / Laboratorio", "Analista de Laboratorio",
     "Dashboard Laboratorio, Consumo Lab.",
     "Resultados microbiol\u00f3gicos, consumo de reactivos, stock insumos",
     "Diaria",
     "PC"],
    ["Calidad / Laboratorio", "Inspector de Calidad",
     "Reportes de Inspecci\u00f3n, Consumos Calidad",
     "Par\u00e1metros de inspecci\u00f3n, no conformidades, consumo suministros",
     "Por lote / turno",
     "Tablet / M\u00f3vil"],
    ["Calidad / Laboratorio", "Jefe de Calidad",
     "Portal completo: Lab, Consumos, Costos, Certificaciones",
     "Aprobaciones, an\u00e1lisis de tendencias, programaci\u00f3n auditor\u00edas",
     "Diaria",
     "PC"],
    ["Almac\u00e9n / Stock", "Almacenero",
     "Stock General (Portal)",
     "Ingresos/salidas de PT, ubicaciones, lotes",
     "Por movimiento",
     "Tablet / M\u00f3vil"],
    ["Almac\u00e9n / Stock", "Jefe de Almac\u00e9n",
     "Dashboard Stock General, Control Temperaturas",
     "Validaci\u00f3n inventarios, alertas temperatura, despachos",
     "Diaria / Continua",
     "PC / Tablet"],
    ["Log\u00edstica", "Coordinador Log\u00edstico",
     "Avance de Contenedores, Stock General",
     "Estado de carga, documentaci\u00f3n de embarque",
     "Por despacho",
     "PC"],
    ["Gerencia", "Gerente General",
     "Portal Gerencia: todos los m\u00f3dulos",
     "Revisi\u00f3n de KPIs, aprobaciones, directivas",
     "Diaria / Semanal",
     "PC / M\u00f3vil"],
    ["Gerencia", "Gerente Comercial",
     "Stock General, Web Corporativa",
     "Consulta disponibilidad, coordinaci\u00f3n con clientes",
     "Diaria",
     "PC / M\u00f3vil"],
]

create_styled_table(workflow_headers, workflow_rows)

add_page_break()

# ══════════════════════════════════════════════════════════════════
# SECTION 5: BENEFICIOS CLAVE
# ══════════════════════════════════════════════════════════════════

add_section_header("5", "Beneficios Clave para la Empresa")

add_formatted_paragraph(
    "La implementaci\u00f3n del ecosistema digital ha generado beneficios tangibles y medibles en "
    "m\u00faltiples dimensiones de la operaci\u00f3n:",
    size=11, space_before=6, space_after=10
)

add_subsection_header("Eficiencia Operativa")
benefits_ops = [
    ("Reducci\u00f3n del 70% en tiempo de generaci\u00f3n de reportes: ",
     "Lo que antes tomaba horas de consolidaci\u00f3n manual en Excel ahora est\u00e1 disponible en tiempo real con un clic."),
    ("Visibilidad inmediata de la producci\u00f3n: ",
     "El tracking hora a hora permite detectar desviaciones en el momento y corregir antes de que impacten el resultado del turno."),
    ("Eliminaci\u00f3n de errores de transcripci\u00f3n: ",
     "Al ingresar datos directamente al sistema, se eliminan los errores de copiar datos de papel a Excel."),
]
for bp, txt in benefits_ops:
    add_bullet(txt, bold_prefix=bp)

add_subsection_header("Control de Calidad y Cumplimiento")
benefits_qa = [
    ("Trazabilidad digital completa: ",
     "Cada dato queda registrado con fecha, hora y usuario, cumpliendo requisitos de BRC, HACCP y FDA."),
    ("Alertas preventivas de temperatura: ",
     "7 alertas activas en tiempo real protegen la cadena de fr\u00edo y previenen p\u00e9rdidas millonarias por producto no conforme."),
    ("Gesti\u00f3n proactiva de insumos de laboratorio: ",
     "El sem\u00e1foro de stock (CR\u00cdTICO/BAJO/MODERADO/OK) para 30+ productos evita paradas por falta de reactivos."),
]
for bp, txt in benefits_qa:
    add_bullet(txt, bold_prefix=bp)

add_subsection_header("Gesti\u00f3n Gerencial")
benefits_mgmt = [
    ("Toma de decisiones basada en datos: ",
     "La gerencia puede tomar decisiones informadas en minutos, no en d\u00edas."),
    ("Control centralizado multi-almac\u00e9n: ",
     "Visibilidad de 272,020 kg de stock en 3 almacenes (FTP, Agroempaques, Ransa) desde un solo dashboard."),
    ("Imagen corporativa profesional: ",
     "La p\u00e1gina web posiciona a Frutos Tropicales como empresa moderna ante clientes internacionales."),
]
for bp, txt in benefits_mgmt:
    add_bullet(txt, bold_prefix=bp)

add_subsection_header("Ahorro Econ\u00f3mico")
benefits_cost = [
    ("Reducci\u00f3n de horas-hombre en reportes: ",
     "Estimado de 15-20 horas semanales ahorradas en generaci\u00f3n manual de informes."),
    ("Prevenci\u00f3n de p\u00e9rdidas por quiebre de cadena de fr\u00edo: ",
     "Un solo evento de temperatura no detectado puede significar p\u00e9rdidas de US$ 10,000 - 50,000 en producto."),
    ("Optimizaci\u00f3n de inventarios: ",
     "Mejor visibilidad de stock reduce sobre-almacenamiento y riesgo de producto vencido."),
    ("Desarrollo interno vs. tercerizado: ",
     "El ecosistema completo fue desarrollado internamente, ahorrando decenas de miles de d\u00f3lares en consultor\u00edas y licencias de software empresarial."),
]
for bp, txt in benefits_cost:
    add_bullet(txt, bold_prefix=bp)

add_page_break()

# ══════════════════════════════════════════════════════════════════
# SECTION 6: MEJORAS PLANIFICADAS A CORTO PLAZO
# ══════════════════════════════════════════════════════════════════

add_section_header("6", "Mejoras Planificadas a Corto Plazo")

add_formatted_paragraph(
    "El ecosistema actual sienta las bases para mejoras incrementales que multiplicar\u00e1n "
    "su valor. Las siguientes mejoras est\u00e1n planificadas para los pr\u00f3ximos 3 a 6 meses:",
    size=11, space_before=6, space_after=10
)

improvements_headers = ["Mejora", "Descripci\u00f3n", "Impacto Esperado", "Plazo"]

improvements_rows = [
    ["Filtros interactivos avanzados",
     "Ampliar los filtros en todos los dashboards: por rango de fechas, turno, l\u00ednea de producci\u00f3n, cliente, destino",
     "Mayor granularidad en an\u00e1lisis, reportes personalizados por necesidad",
     "1-2 meses"],
    ["Vistas optimizadas para m\u00f3vil",
     "Redise\u00f1o responsive espec\u00edfico para smartphones y tablets de campo",
     "Supervisores e inspectores acceden a datos sin ir a oficina",
     "2-3 meses"],
    ["Alertas autom\u00e1ticas por Email y WhatsApp",
     "Notificaciones autom\u00e1ticas cuando temperaturas, stock o producci\u00f3n salen de rango",
     "Respuesta inmediata a eventos cr\u00edticos, 24/7 sin monitoreo manual",
     "2-3 meses"],
    ["Integraci\u00f3n con sistema ERP",
     "Conexi\u00f3n bidireccional con el sistema contable/ERP existente",
     "Eliminaci\u00f3n de doble digitaci\u00f3n, datos financieros integrados",
     "3-6 meses"],
    ["Reportes PDF autom\u00e1ticos",
     "Generaci\u00f3n programada de reportes en PDF para distribuci\u00f3n por correo",
     "Gerencia recibe resumen diario autom\u00e1tico cada ma\u00f1ana",
     "1-2 meses"],
    ["Historial y tendencias",
     "Almacenamiento hist\u00f3rico para comparar semanas, meses, campa\u00f1as",
     "Identificar patrones estacionales, mejorar planificaci\u00f3n",
     "2-4 meses"],
    ["Dashboard de KPIs financieros",
     "Costo por kg procesado, costo por contenedor, margen por producto",
     "Visibilidad directa del impacto financiero de la operaci\u00f3n",
     "3-4 meses"],
]

create_styled_table(improvements_headers, improvements_rows)

add_page_break()

# ══════════════════════════════════════════════════════════════════
# SECTION 7: VISION A FUTURO - PLANTA DE FRESCOS
# ══════════════════════════════════════════════════════════════════

add_section_header("7", "Visi\u00f3n a Futuro \u2014 Planta de Frescos")

add_formatted_paragraph(
    "La empresa cuenta con una planta de frutas frescas que actualmente opera con procesos "
    "mayormente manuales. Extender el ecosistema digital a esta planta representa una oportunidad "
    "estrat\u00e9gica para estandarizar operaciones y obtener visibilidad integral del negocio.",
    size=11, space_before=6, space_after=10
)

add_subsection_header("Nuevos m\u00f3dulos propuestos para Planta de Frescos:")

fresh_modules = [
    ("Dashboard de Recepci\u00f3n y Clasificaci\u00f3n: ",
     "Registro digital de materia prima recibida: proveedor, variedad, peso, calibre, "
     "estado de madurez, procedencia. Clasificaci\u00f3n autom\u00e1tica por categor\u00eda de calidad."),
    ("Control de Calidad de Fruta Fresca: ",
     "Dashboard espec\u00edfico para par\u00e1metros de calidad de fruta fresca: grados Brix, firmeza, "
     "color, defectos visuales, porcentaje de descarte. Sem\u00e1foro de aceptaci\u00f3n/rechazo."),
    ("Monitoreo de L\u00edneas de Empaque: ",
     "Seguimiento en tiempo real de cada l\u00ednea de empaque: velocidad, merma, cajas producidas, "
     "eficiencia por operador, paradas programadas y no programadas."),
    ("Trazabilidad de Cadena de Fr\u00edo (Campo a Cliente): ",
     "Registro continuo de temperatura desde cosecha, transporte, recepci\u00f3n, almacenamiento, "
     "empaque y despacho. Certificaci\u00f3n digital de cadena de fr\u00edo para cada lote."),
    ("Control de Peso y Calibre: ",
     "Dashboard para monitorear calibre y peso de fruta por lote, proveedor y variedad. "
     "Alertas cuando el producto no cumple especificaciones del cliente."),
    ("Dashboards de Empaque por Cliente: ",
     "Paneles personalizados por cliente con especificaciones de empaque, etiquetado, "
     "calibres requeridos, y avance de pedidos en tiempo real."),
    ("Integraci\u00f3n con Sistema de Congelados: ",
     "Unificaci\u00f3n de datos de ambas plantas en un portal gerencial \u00fanico. "
     "Comparativos de rendimiento, costos y productividad entre plantas."),
]

for bp, txt in fresh_modules:
    add_bullet(txt, bold_prefix=bp)

add_formatted_paragraph(
    "La implementaci\u00f3n de estos m\u00f3dulos permitir\u00e1 a la gerencia tener una visi\u00f3n 360\u00b0 de "
    "ambas operaciones desde un solo portal, facilitando la toma de decisiones estrat\u00e9gicas "
    "y la optimizaci\u00f3n de recursos compartidos.",
    size=11, color=GREEN_DARK, space_before=10, space_after=6, bold=True
)

add_page_break()

# ══════════════════════════════════════════════════════════════════
# SECTION 8: VISION A FUTURO - INNOVACIONES TECNOLOGICAS
# ══════════════════════════════════════════════════════════════════

add_section_header("8", "Visi\u00f3n a Futuro \u2014 Innovaciones Tecnol\u00f3gicas")

add_formatted_paragraph(
    "El ecosistema actual puede evolucionar incorporando tecnolog\u00edas emergentes que "
    "potenciar\u00e1n a\u00fan m\u00e1s la competitividad de Frutos Tropicales Peru:",
    size=11, space_before=6, space_after=10
)

innovations_headers = ["Innovaci\u00f3n", "Descripci\u00f3n", "Beneficio para FTP"]

innovations_rows = [
    ["Sensores IoT",
     "Sensores autom\u00e1ticos de temperatura y humedad en c\u00e1maras, transporte y almacenes con transmisi\u00f3n en tiempo real al dashboard",
     "Elimina lectura manual, monitoreo 24/7 sin intervenci\u00f3n humana, historial autom\u00e1tico"],
    ["Aplicaci\u00f3n M\u00f3vil Nativa",
     "App dedicada para supervisores e inspectores con funcionalidad offline, c\u00e1mara para evidencia fotogr\u00e1fica y notificaciones push",
     "Acceso sin conexi\u00f3n en campo, registro fotogr\u00e1fico de no conformidades, alertas instant\u00e1neas"],
    ["Anal\u00edtica Predictiva con IA",
     "Modelos de machine learning para predecir rendimiento, calidad de materia prima, demanda y mantenimiento de equipos",
     "Anticipar problemas antes de que ocurran, optimizar planificaci\u00f3n de producci\u00f3n"],
    ["Generaci\u00f3n Autom\u00e1tica de Reportes PDF",
     "Sistema programado que genera y distribuye reportes personalizados por correo cada d\u00eda, semana o mes",
     "Gerencia recibe informaci\u00f3n sin tener que buscarla, formato profesional para clientes"],
    ["Business Intelligence (BI)",
     "Dashboards avanzados con drill-down, an\u00e1lisis multidimensional y modelos de datos complejos",
     "An\u00e1lisis estrat\u00e9gico profundo, identificar oportunidades de mejora ocultas en los datos"],
    ["Integraci\u00f3n Log\u00edstica",
     "Conexi\u00f3n con plataformas de naviera, agentes de aduana y booking de contenedores",
     "Trazabilidad end-to-end, reducci\u00f3n de tiempos de coordinaci\u00f3n, visibilidad de embarques"],
    ["Blockchain para Trazabilidad",
     "Registro inmutable de trazabilidad desde campo hasta cliente final, verificable por compradores",
     "M\u00e1xima transparencia y confianza, diferenciador competitivo en mercados premium"],
]

create_styled_table(innovations_headers, innovations_rows)

add_page_break()

# ══════════════════════════════════════════════════════════════════
# SECTION 9: JUSTIFICACION DE INVERSION - PLAN MAX
# ══════════════════════════════════════════════════════════════════

add_section_header("9", "Justificaci\u00f3n de la Inversi\u00f3n \u2014 Plan MAX")

add_formatted_paragraph(
    "Para que el ecosistema digital alcance su m\u00e1ximo potencial y soporte las mejoras planificadas, "
    "las innovaciones tecnol\u00f3gicas y la expansi\u00f3n a la planta de frescos, es indispensable contar "
    "con el Plan MAX de suscripci\u00f3n, el plan m\u00e1s completo disponible.",
    size=11, space_before=6, space_after=10
)

add_subsection_header("\u00bfPor qu\u00e9 el Plan MAX y no un plan b\u00e1sico o gratuito?")

max_reasons = [
    ("Capacidad de procesamiento y almacenamiento: ",
     "El ecosistema maneja datos en tiempo real de m\u00faltiples fuentes (producci\u00f3n, temperaturas, "
     "laboratorio, stock). Los planes b\u00e1sicos o gratuitos tienen l\u00edmites que impedir\u00edan el "
     "funcionamiento simult\u00e1neo de todos los dashboards y la cantidad de consultas necesarias."),
    ("Velocidad y rendimiento: ",
     "El Plan MAX garantiza tiempos de respuesta r\u00e1pidos incluso con m\u00faltiples usuarios "
     "accediendo simult\u00e1neamente. En una operaci\u00f3n que procesa 14+ toneladas diarias, "
     "cada segundo cuenta."),
    ("Funcionalidades avanzadas: ",
     "Alertas autom\u00e1ticas, integraciones con APIs externas, almacenamiento de historial extendido, "
     "y capacidad de procesamiento para anal\u00edtica avanzada solo est\u00e1n disponibles en el Plan MAX."),
    ("Soporte prioritario y SLA: ",
     "Un sistema cr\u00edtico para la operaci\u00f3n necesita soporte garantizado. El Plan MAX incluye "
     "tiempos de respuesta comprometidos y asistencia t\u00e9cnica prioritaria."),
    ("Escalabilidad para la Planta de Frescos: ",
     "La expansi\u00f3n a la planta de frescos duplicar\u00e1 la cantidad de datos y usuarios. "
     "Solo el Plan MAX soporta esta escalabilidad sin degradar el rendimiento."),
    ("Seguridad empresarial: ",
     "El Plan MAX incluye cifrado avanzado, backups autom\u00e1ticos, control de acceso "
     "granular y cumplimiento con est\u00e1ndares de seguridad exigidos por nuestras certificaciones."),
]

for bp, txt in max_reasons:
    add_bullet(txt, bold_prefix=bp)

add_subsection_header("An\u00e1lisis Costo-Beneficio")

add_formatted_paragraph(
    "A continuaci\u00f3n se presenta un an\u00e1lisis comparativo que demuestra que la inversi\u00f3n "
    "en el Plan MAX se recupera en pocas semanas:",
    size=11, space_before=4, space_after=8
)

roi_headers = ["Concepto", "Sin Ecosistema Digital", "Con Ecosistema + Plan MAX"]

roi_rows = [
    ["Tiempo en generaci\u00f3n de reportes",
     "15-20 horas/semana (manual en Excel)",
     "Autom\u00e1tico, en tiempo real"],
    ["Riesgo de p\u00e9rdida por temperatura",
     "Alto (detecci\u00f3n manual, tard\u00eda)",
     "M\u00ednimo (alertas autom\u00e1ticas 24/7)"],
    ["Costo de un quiebre de cadena de fr\u00edo",
     "US$ 10,000 - 50,000 por evento",
     "Prevenido con monitoreo continuo"],
    ["Visibilidad de stock en tiempo real",
     "Aproximada, con d\u00edas de retraso",
     "Exacta, al instante, multi-almac\u00e9n"],
    ["Trazabilidad para auditor\u00edas",
     "Parcial, basada en registros en papel",
     "Completa, digital, con timestamps"],
    ["Costo estimado de desarrollo externo equivalente",
     "US$ 50,000 - 100,000+ (consultor\u00eda)",
     "Desarrollado internamente por Rodrigo"],
    ["Tiempo de despliegue (consultor\u00eda externa)",
     "6-12 meses m\u00ednimo",
     "Ya implementado y operativo"],
    ["Capacitaci\u00f3n del equipo",
     "Compleja (sistema externo desconocido)",
     "R\u00e1pida (desarrollador interno conoce la operaci\u00f3n)"],
]

create_styled_table(roi_headers, roi_rows)

add_formatted_paragraph("")

add_subsection_header("Retorno de la Inversi\u00f3n (ROI)")

roi_items = [
    ("Prevenci\u00f3n de 1 solo evento de temperatura", "= ahorro de US$ 10,000 - 50,000"),
    ("Ahorro en horas-hombre de reportes", "= ~US$ 500-800/mes"),
    ("Optimizaci\u00f3n de inventarios y reducci\u00f3n de merma", "= US$ 2,000-5,000/mes estimado"),
    ("Ahorro vs. desarrollo externo equivalente", "= US$ 50,000-100,000 (costo evitado)"),
    ("Costo del Plan MAX", "= una fracci\u00f3n m\u00ednima de los ahorros generados"),
]

for label, value in roi_items:
    add_kpi_box(label, value)

add_formatted_paragraph(
    "La inversi\u00f3n en el Plan MAX se paga sola con la prevenci\u00f3n de un solo incidente "
    "de temperatura o con los ahorros mensuales en eficiencia operativa. No invertir en esta "
    "herramienta cr\u00edtica es asumir un riesgo innecesario para la operaci\u00f3n.",
    bold=True, size=11, color=GREEN_DARK, space_before=12, space_after=6
)

add_page_break()

# ══════════════════════════════════════════════════════════════════
# SECTION 10: CONCLUSION
# ══════════════════════════════════════════════════════════════════

add_section_header("10", "Conclusi\u00f3n")

add_formatted_paragraph(
    "Frutos Tropicales Peru S.A.C. cuenta hoy con un ecosistema digital de gesti\u00f3n integral "
    "que pocas empresas de nuestro sector poseen. Este ecosistema, desarrollado \u00edntegramente por "
    "Rodrigo Garcia, ha transformado la manera en que la empresa gestiona su informaci\u00f3n operativa, "
    "de calidad, de laboratorio y gerencial.",
    size=11, space_before=6, space_after=10
)

add_formatted_paragraph(
    "Los 15+ componentes del ecosistema \u2014 desde la p\u00e1gina web corporativa hasta los dashboards "
    "interactivos de stock, producci\u00f3n, temperaturas y laboratorio \u2014 representan una ventaja "
    "competitiva real y medible. La informaci\u00f3n que antes tardaba d\u00edas en consolidarse ahora "
    "est\u00e1 disponible en tiempo real, las alertas de temperatura protegen millones en inventario, "
    "y la trazabilidad digital fortalece nuestras certificaciones internacionales.",
    size=11, space_before=4, space_after=10
)

add_formatted_paragraph(
    "La visi\u00f3n a futuro es clara: extender este ecosistema a la planta de frescos, incorporar "
    "tecnolog\u00edas IoT y anal\u00edtica predictiva, y consolidar una plataforma que cubra el 100% "
    "de las operaciones de la empresa. Para hacer realidad esta visi\u00f3n, es necesario contar "
    "con las herramientas adecuadas.",
    size=11, space_before=4, space_after=10
)

# Call to action box with green background
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(16)
p.paragraph_format.space_after = Pt(4)
pPr = p._p.get_or_add_pPr()
pBdr = parse_xml(
    f'<w:pBdr {nsdecls("w")}>'
    f'<w:top w:val="single" w:sz="12" w:space="6" w:color="16a34a"/>'
    f'<w:left w:val="single" w:sz="12" w:space="6" w:color="16a34a"/>'
    f'<w:bottom w:val="single" w:sz="12" w:space="6" w:color="16a34a"/>'
    f'<w:right w:val="single" w:sz="12" w:space="6" w:color="16a34a"/>'
    f'</w:pBdr>'
)
_insert_pPr_child(pPr, pBdr)
# Add shading to the paragraph
shading = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="f0fdf4"/>')
_insert_pPr_child(pPr, shading)

run = p.add_run("SOLICITUD A GERENCIA GENERAL")
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = GREEN_DARK

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_before = Pt(0)
p2.paragraph_format.space_after = Pt(4)
pPr2 = p2._p.get_or_add_pPr()
pBdr2 = parse_xml(
    f'<w:pBdr {nsdecls("w")}>'
    f'<w:left w:val="single" w:sz="12" w:space="6" w:color="16a34a"/>'
    f'<w:right w:val="single" w:sz="12" w:space="6" w:color="16a34a"/>'
    f'</w:pBdr>'
)
_insert_pPr_child(pPr2, pBdr2)
shading2 = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="f0fdf4"/>')
_insert_pPr_child(pPr2, shading2)

run2 = p2.add_run(
    "Se solicita la aprobaci\u00f3n del Plan MAX de suscripci\u00f3n para garantizar la continuidad, "
    "escalabilidad y evoluci\u00f3n del ecosistema digital que ya est\u00e1 generando valor para la empresa."
)
run2.font.size = Pt(11)
run2.font.color.rgb = GREEN_DARK

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_before = Pt(0)
p3.paragraph_format.space_after = Pt(16)
pPr3 = p3._p.get_or_add_pPr()
pBdr3 = parse_xml(
    f'<w:pBdr {nsdecls("w")}>'
    f'<w:left w:val="single" w:sz="12" w:space="6" w:color="16a34a"/>'
    f'<w:bottom w:val="single" w:sz="12" w:space="6" w:color="16a34a"/>'
    f'<w:right w:val="single" w:sz="12" w:space="6" w:color="16a34a"/>'
    f'</w:pBdr>'
)
_insert_pPr_child(pPr3, pBdr3)
shading3 = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="f0fdf4"/>')
_insert_pPr_child(pPr3, shading3)

run3 = p3.add_run(
    "La inversi\u00f3n es m\u00ednima comparada con el valor que ya genera y el potencial de crecimiento "
    "que habilita. Cada d\u00eda sin estas herramientas completas es un d\u00eda de oportunidades perdidas."
)
run3.bold = True
run3.font.size = Pt(11)
run3.font.color.rgb = GREEN_PRIMARY

# Signature area
add_formatted_paragraph("", size=24)

p_sign = doc.add_paragraph()
p_sign.alignment = WD_ALIGN_PARAGRAPH.CENTER
pPr_s = p_sign._p.get_or_add_pPr()
pBdr_s = parse_xml(
    f'<w:pBdr {nsdecls("w")}>'
    f'<w:top w:val="single" w:sz="4" w:space="8" w:color="16a34a"/>'
    f'</w:pBdr>'
)
_insert_pPr_child(pPr_s, pBdr_s)
run_s = p_sign.add_run("Rodrigo Garcia")
run_s.bold = True
run_s.font.size = Pt(13)
run_s.font.color.rgb = GREEN_DARK

add_formatted_paragraph("Systems Analyst & Developer-QA", bold=False, size=11,
                        color=GREEN_PRIMARY, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                        space_before=0, space_after=2, italic=True)

add_formatted_paragraph("Frutos Tropicales Peru S.A.C.", bold=False, size=10,
                        color=GRAY, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                        space_before=0, space_after=2)

add_formatted_paragraph("20 de marzo de 2026", bold=False, size=10,
                        color=GRAY, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                        space_before=0, space_after=0)

# ══════════════════════════════════════════════════════════════════
# HEADERS AND FOOTERS
# ══════════════════════════════════════════════════════════════════

for section in doc.sections:
    # Header
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run("Frutos Tropicales Peru S.A.C. | Informe Ejecutivo \u2014 Ecosistema Digital")
    hr.font.size = Pt(8)
    hr.font.color.rgb = GREEN_PRIMARY
    hr.italic = True

    # Footer with page numbers
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add "CONFIDENCIAL" text
    fr = fp.add_run("DOCUMENTO CONFIDENCIAL \u2014 Frutos Tropicales Peru S.A.C. \u2014 P\u00e1gina ")
    fr.font.size = Pt(7)
    fr.font.color.rgb = GRAY_LIGHT

    # Page number field
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')

    run_pg = fp.add_run()
    run_pg.font.size = Pt(7)
    run_pg.font.color.rgb = GRAY_LIGHT
    run_pg._r.append(fldChar1)
    run_pg._r.append(instrText)
    run_pg._r.append(fldChar2)

# ══════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════

# Fix settings.xml zoom percent attribute (required by OOXML schema)
settings = doc.settings.element
zoom_elements = settings.findall(qn('w:zoom'))
for z in zoom_elements:
    if z.get(qn('w:percent')) is None:
        z.set(qn('w:percent'), '100')

# If no zoom element exists, add one
if not zoom_elements:
    zoom_el = parse_xml(f'<w:zoom {nsdecls("w")} w:percent="100"/>')
    settings.append(zoom_el)

doc.save(OUTPUT_PATH)
print(f"Documento guardado exitosamente en: {OUTPUT_PATH}")
print(f"Tamano: {os.path.getsize(OUTPUT_PATH):,} bytes")
